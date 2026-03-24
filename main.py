from fastapi import FastAPI, Request
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse, FileResponse
from fastapi.templating import Jinja2Templates
import uvicorn
import logging
import uuid
import tempfile
from typing import Optional
from datetime import datetime
from database import db
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="WebCinema")

# Настройка шаблонов
templates = Jinja2Templates(directory="templates")

# Сессии (простое хранилище)
sessions = {}


def get_current_user(request: Request):
    session_id = request.cookies.get("session_id")
    if session_id and session_id in sessions:
        return sessions[session_id]
    return None


# ========== СТРАНИЦЫ ==========

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.get("/schedule", response_class=HTMLResponse)
async def schedule(request: Request):
    return templates.TemplateResponse("schedule.html", {"request": request})

@app.get("/all-films", response_class=HTMLResponse)
async def all_films(request: Request):
    return templates.TemplateResponse("all-films.html", {"request": request})

@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    return templates.TemplateResponse("login.html", {"request": request})


@app.get("/admin", response_class=HTMLResponse)
async def admin_panel(request: Request):
    user = get_current_user(request)
    if not user or user['role'] != 'admin':
        return RedirectResponse(url="/login")
    return templates.TemplateResponse("admin/panel.html", {"request": request})


@app.get("/admin/movies", response_class=HTMLResponse)
async def admin_movies(request: Request):
    user = get_current_user(request)
    if not user or user['role'] != 'admin':
        return RedirectResponse(url="/login")
    return templates.TemplateResponse("admin/movies.html", {"request": request})


@app.get("/admin/sessions", response_class=HTMLResponse)
async def admin_sessions(request: Request):
    user = get_current_user(request)
    if not user or user['role'] != 'admin':
        return RedirectResponse(url="/login")
    return templates.TemplateResponse("admin/sessions.html", {"request": request})


@app.get("/admin/tariffs", response_class=HTMLResponse)
async def admin_tariffs(request: Request):
    user = get_current_user(request)
    if not user or user['role'] != 'admin':
        return RedirectResponse(url="/login")
    return templates.TemplateResponse("admin/tariffs.html", {"request": request})


@app.get("/admin/users", response_class=HTMLResponse)
async def admin_users(request: Request):
    user = get_current_user(request)
    if not user or user['role'] != 'admin':
        return RedirectResponse(url="/login")
    return templates.TemplateResponse("admin/users.html", {"request": request})


@app.get("/cashier", response_class=HTMLResponse)
async def cashier_panel(request: Request):
    user = get_current_user(request)
    if not user or user['role'] not in ['admin', 'cashier']:
        return RedirectResponse(url="/login")
    return templates.TemplateResponse("cashier/panel.html", {"request": request})


@app.get("/cashier/sales", response_class=HTMLResponse)
async def cashier_sales(request: Request):
    user = get_current_user(request)
    if not user or user['role'] not in ['admin', 'cashier']:
        return RedirectResponse(url="/login")
    return templates.TemplateResponse("cashier/sales.html", {"request": request})


@app.get("/cashier/hall/{session_id}", response_class=HTMLResponse)
async def cashier_hall(request: Request, session_id: int):
    user = get_current_user(request)
    if not user or user['role'] not in ['admin', 'cashier']:
        return RedirectResponse(url="/login")
    return templates.TemplateResponse("cashier/hall.html", {
        "request": request,
        "session_id": session_id
    })


@app.get("/cashier/checkout")
async def cashier_checkout(
    request: Request,
    session_id: int,
    seat_id: int,
    row: int,
    seat: int,
    price: int
):
    try:
        session = db.get_session_by_id(session_id)
        if not session:
            return HTMLResponse("Сеанс не найден", status_code=404)

        return templates.TemplateResponse(
            "cashier/checkout.html",
            {
                "request": request,
                "session": session,
                "seat_id": seat_id,
                "row": row,
                "seat_number": seat,
                "price": price
            }
        )
    except Exception as e:
        logger.error(f"Ошибка: {e}")
        return HTMLResponse(f"Ошибка: {e}", status_code=500)


@app.get("/cashier/returns", response_class=HTMLResponse)
async def cashier_returns(request: Request):
    user = get_current_user(request)
    if not user or user['role'] not in ['admin', 'cashier']:
        return RedirectResponse(url="/login")
    return templates.TemplateResponse("cashier/returns.html", {"request": request})


@app.get("/cashier/tickets", response_class=HTMLResponse)
async def cashier_tickets(request: Request):
    return templates.TemplateResponse("cashier/tickets.html", {"request": request})


# ========== API ФИЛЬМЫ ==========

@app.get("/api/movies")
async def get_movies():
    try:
        movies = db.get_movies()
        return JSONResponse(content=movies)
    except Exception as e:
        logger.error(f"Ошибка: {e}")
        return JSONResponse(content=[], status_code=500)


@app.post("/api/movies/add")
async def add_movie(request: Request):
    try:
        data = await request.json()
        movie_id = db.add_movie(
            data['title'], data['duration'], data['genre'], data['age_rating']
        )
        return {"success": True, "id": movie_id}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/movies/update/{movie_id}")
async def update_movie(request: Request, movie_id: int):
    try:
        data = await request.json()
        query = """
            UPDATE movies 
            SET title = %s, duration = %s, genre = %s, age_rating = %s
            WHERE id = %s
        """
        result = db.execute_query(query, (
            data['title'], data['duration'], data['genre'], data['age_rating'], movie_id
        ))
        return {"success": result > 0}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/movies/delete/{movie_id}")
async def delete_movie(request: Request, movie_id: int):
    try:
        user = get_current_user(request)
        user_id = user['id'] if user else None
        success = db.delete_movie(movie_id, deleted_by=user_id)
        return {"success": success}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ========== API СЕАНСЫ ==========

@app.get("/api/sessions")
async def get_sessions():
    try:
        sessions = db.get_sessions()
        return JSONResponse(content=sessions)
    except Exception as e:
        return JSONResponse(content=[], status_code=500)


@app.get("/api/session/{session_id}")
async def get_session(session_id: int):
    try:
        session = db.get_session_by_id(session_id)
        return JSONResponse(content=session)
    except Exception as e:
        return JSONResponse(content=None, status_code=500)


@app.post("/api/sessions/add")
async def add_session(request: Request):
    try:
        data = await request.json()
        result = db.add_session(
            data['session_date'], data['movie_id'],
            data['start_time'], data['end_time'], data['hall_id']
        )
        return result
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/sessions/update/{session_id}")
async def update_session(request: Request, session_id: int):
    try:
        data = await request.json()
        result = db.update_session(
            session_id, data['session_date'], data['movie_id'],
            data['start_time'], data['end_time'], data['hall_id']
        )
        return result
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/sessions/delete/{session_id}")
async def delete_session(request: Request, session_id: int):
    try:
        success = db.delete_session(session_id)
        return {"success": success}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ========== API ТАРИФЫ ==========

@app.get("/api/tariffs")
async def get_tariffs():
    try:
        tariffs = db.get_tariffs()
        return JSONResponse(content=tariffs)
    except Exception as e:
        return JSONResponse(content=[], status_code=500)


@app.post("/api/tariffs/add")
async def add_tariff(request: Request):
    try:
        data = await request.json()
        tariff_id = db.add_tariff(data['name'], data['price'])
        return {"success": True, "id": tariff_id}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/tariffs/update/{tariff_id}")
async def update_tariff(request: Request, tariff_id: int):
    try:
        data = await request.json()
        success = db.update_tariff(tariff_id, data['name'], data['price'])
        return {"success": success}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/tariffs/delete/{tariff_id}")
async def delete_tariff(request: Request, tariff_id: int):
    try:
        success = db.delete_tariff(tariff_id)
        return {"success": success}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ========== API МЕСТА ==========

@app.get("/api/seats/{hall_id}")
async def get_seats(hall_id: int):
    try:
        seats = db.get_seats(hall_id)
        return JSONResponse(content=seats)
    except Exception as e:
        return JSONResponse(content=[], status_code=500)


# ========== API БИЛЕТЫ ==========

@app.get("/api/tickets-list")
async def get_all_tickets(
    session_id: Optional[int] = None,
    date: Optional[str] = None
):
    """Получение всех билетов с фильтрацией"""
    try:
        query = """
            SELECT 
                t.id,
                t.customer_name,
                t.price,
                t.payment_method,
                t.is_returned,
                to_char(s.session_date, 'DD.MM.YYYY') as session_date,
                to_char(s.start_time, 'HH24:MI') as start_time,
                to_char(s.end_time, 'HH24:MI') as end_time,
                m.title as movie,
                h.hall_number as hall,
                se.row_number,
                se.seat_number
            FROM tickets t
            LEFT JOIN sessions s ON t.session_id = s.id
            LEFT JOIN movies m ON s.movie_id = m.id
            LEFT JOIN halls h ON s.hall_id = h.id
            LEFT JOIN seats se ON t.seat_id = se.id
            WHERE 1=1
        """
        params = []

        if session_id:
            query += " AND s.id = %s"
            params.append(session_id)

        if date:
            query += " AND s.session_date = %s"
            params.append(date)

        query += " ORDER BY t.sold_at DESC"

        tickets = db.execute_query(query, tuple(params) if params else None, fetch_all=True)

        if tickets is None:
            tickets = []

        for ticket in tickets:
            if ticket.get('session_date') and hasattr(ticket['session_date'], 'strftime'):
                ticket['session_date'] = ticket['session_date'].strftime('%d.%m.%Y')
            if ticket.get('start_time') and hasattr(ticket['start_time'], 'strftime'):
                ticket['start_time'] = ticket['start_time'].strftime('%H:%M')
            if ticket.get('end_time') and hasattr(ticket['end_time'], 'strftime'):
                ticket['end_time'] = ticket['end_time'].strftime('%H:%M')

        return JSONResponse(content=tickets)

    except Exception as e:
        logger.error(f"Ошибка: {e}")
        return JSONResponse(content=[], status_code=500)


@app.get("/api/tickets/{session_id}")
async def get_tickets(session_id: int):
    """Получение проданных билетов"""
    try:
        tickets = db.get_sold_tickets(session_id)
        logger.info(f"Возвращаем билеты для сеанса {session_id}: {len(tickets)} шт.")
        return JSONResponse(content=tickets)
    except Exception as e:
        logger.error(f"Ошибка: {e}")
        return JSONResponse(content=[], status_code=500)


@app.post("/api/tickets/buy")
async def buy_ticket(request: Request):
    """Покупка билета"""
    try:
        data = await request.json()
        logger.info(f"Покупка билета: {data}")

        # Проверяем, не продан ли билет
        check_query = """
            SELECT id FROM tickets 
            WHERE session_id = %s AND seat_id = %s AND is_returned = false
        """
        existing = db.execute_query(
            check_query, (data['session_id'], data['seat_id']), fetch_one=True
        )

        if existing:
            return {"success": False, "error": "Это место уже занято"}

        # Продаем билет
        insert_query = """
            INSERT INTO tickets (session_id, seat_id, customer_name, price, payment_method, sold_at, is_returned)
            VALUES (%s, %s, %s, %s, %s, NOW(), false)
            RETURNING id
        """
        result = db.execute_query(
            insert_query,
            (data['session_id'], data['seat_id'], data['customer_name'],
             data['price'], data['payment_method']),
            fetch_one=True
        )

        logger.info(f"Результат вставки: {result}")

        if result and result.get('id'):
            ticket_id = result['id']
            logger.info(f"Билет продан, ID: {ticket_id}")
            return {"success": True, "id": ticket_id}
        else:
            return {"success": False, "error": "Не удалось сохранить билет"}

    except Exception as e:
        logger.error(f"Ошибка: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return {"success": False, "error": str(e)}


@app.post("/api/tickets/return/{ticket_id}")
async def return_ticket(request: Request, ticket_id: int):
    try:
        user = get_current_user(request)
        user_id = user['id'] if user else None
        success = db.return_ticket(ticket_id, returned_by=user_id)
        return {"success": success}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ========== API АВТОРИЗАЦИЯ ==========

@app.post("/api/login")
async def login(request: Request):
    try:
        data = await request.json()
        user = db.authenticate(data['username'], data['password'])

        if user:
            session_id = str(uuid.uuid4())
            sessions[session_id] = dict(user)

            response = JSONResponse({
                "success": True,
                "user": {
                    "id": user['id'],
                    "username": user['username'],
                    "role": user['role'],
                    "full_name": user['full_name']
                }
            })
            response.set_cookie(key="session_id", value=session_id)
            return response

        return JSONResponse({"success": False, "error": "Неверный логин или пароль"})

    except Exception as e:
        return JSONResponse({"success": False, "error": str(e)})


@app.get("/api/logout")
async def logout(request: Request):
    session_id = request.cookies.get("session_id")
    if session_id in sessions:
        del sessions[session_id]
    response = RedirectResponse(url="/")
    response.delete_cookie("session_id")
    return response


# ========== API ПОЛЬЗОВАТЕЛИ ==========

@app.get("/api/users")
async def get_users():
    try:
        users = db.get_users()
        return JSONResponse(content=users)
    except Exception as e:
        return JSONResponse(content=[], status_code=500)


@app.post("/api/users/add")
async def add_user(request: Request):
    """Добавление пользователя"""
    try:
        data = await request.json()
        logger.info(f"Добавление пользователя: {data}")

        # Проверяем, не существует ли уже такой логин
        check_query = "SELECT id FROM users WHERE username = %s"
        existing = db.execute_query(check_query, (data['username'],), fetch_one=True)

        if existing:
            return {"success": False, "error": "Пользователь с таким логином уже существует"}

        # Добавляем пользователя в БД
        query = """
            INSERT INTO users (username, password, role, full_name) 
            VALUES (%s, %s, %s, %s) 
            RETURNING id
        """
        result = db.execute_query(
            query,
            (data['username'], data['password'], data['role'], data.get('full_name', '')),
            fetch_one=True
        )

        if result and result.get('id'):
            logger.info(f"Пользователь добавлен с ID: {result['id']}")
            return {"success": True, "id": result['id']}
        else:
            return {"success": False, "error": "Не удалось добавить пользователя"}

    except Exception as e:
        logger.error(f"Ошибка: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return {"success": False, "error": str(e)}


@app.post("/api/users/update/{user_id}")
async def update_user(request: Request, user_id: int):
    try:
        data = await request.json()
        success = db.update_user(
            user_id, data['username'], data['role'],
            data.get('full_name', ''),
            data.get('password') if data.get('password') else None
        )
        return {"success": success}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/api/users/delete/{user_id}")
async def delete_user(request: Request, user_id: int):
    try:
        current_user = get_current_user(request)
        current_user_id = current_user['id'] if current_user else None

        if current_user_id == user_id:
            return {"success": False, "error": "Нельзя удалить самого себя"}

        success = db.delete_user(user_id, deleted_by=current_user_id)
        return {"success": success}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ========== API ОТЧЕТЫ ==========

@app.get("/api/reports/word/{session_id}")
async def generate_word_report(session_id: int):
    try:
        session = db.get_session_by_id(session_id)
        if not session:
            return JSONResponse({"error": "Сеанс не найден"}, status_code=404)

        tickets = db.get_sold_tickets(session_id)

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        title = doc.add_heading('ОТЧЕТ ПО СЕАНСУ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('Информация о сеансе:', level=1)
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Light Grid Accent 1'

        info_table.rows[0].cells[0].text = 'Фильм:'
        info_table.rows[0].cells[1].text = session['movie']
        info_table.rows[1].cells[0].text = 'Дата:'
        info_table.rows[1].cells[1].text = session['date']
        info_table.rows[2].cells[0].text = 'Время:'
        info_table.rows[2].cells[1].text = f"{session['start_time']} - {session['end_time']}"
        info_table.rows[3].cells[0].text = 'Зал:'
        info_table.rows[3].cells[1].text = str(session['hall'])
        info_table.rows[4].cells[0].text = 'Всего билетов:'
        info_table.rows[4].cells[1].text = str(len(tickets))

        doc.add_paragraph()
        doc.add_heading('Проданные билеты:', level=1)

        if tickets:
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Light Grid Accent 1'

            headers = ['№', 'Ряд', 'Место', 'Зритель', 'Цена', 'Оплата', 'Статус']
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header

            for i, ticket in enumerate(tickets, 1):
                row = table.add_row().cells
                row[0].text = str(i)
                row[1].text = str(ticket.get('row_number', ''))
                row[2].text = str(ticket.get('seat_number', ''))
                row[3].text = ticket.get('customer_name', '')
                row[4].text = f"{ticket.get('price', 0)} ₽"
                row[5].text = ticket.get('payment_method', '')
                row[6].text = 'Продан' if not ticket.get('is_returned') else 'Возврат'

            total_sum = sum(t.get('price', 0) for t in tickets if not t.get('is_returned'))
            summary = doc.add_paragraph()
            summary.add_run(f'ИТОГО: {len(tickets)} билетов на сумму {total_sum} ₽').bold = True
        else:
            doc.add_paragraph('Нет проданных билетов')

        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph('Директор кинотеатра: ____________________')
        doc.add_paragraph(f'Дата составления отчета: {datetime.now().strftime("%d.%m.%Y")}')

        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)

        return FileResponse(
            temp_file.name,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=f'отчет_сеанс_{session_id}_{session["date"]}.docx'
        )

    except Exception as e:
        logger.error(f"Ошибка: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)


# ========== БИЛЕТ HTML ==========

@app.get("/ticket/{ticket_id}")
async def get_ticket(ticket_id: int):
    try:
        query = """
            SELECT 
                t.id, t.customer_name, t.price, t.payment_method,
                s.session_date, s.start_time, s.end_time,
                m.title as movie, h.hall_number as hall,
                se.row_number, se.seat_number
            FROM tickets t
            JOIN sessions s ON t.session_id = s.id
            JOIN movies m ON s.movie_id = m.id
            JOIN halls h ON s.hall_id = h.id
            JOIN seats se ON t.seat_id = se.id
            WHERE t.id = %s
        """
        ticket = db.execute_query(query, (ticket_id,), fetch_one=True)

        if not ticket:
            return HTMLResponse("Билет не найден", status_code=404)

        date_str = ticket['session_date'].strftime('%d.%m.%Y')
        time_str = f"{ticket['start_time'].strftime('%H:%M')} - {ticket['end_time'].strftime('%H:%M')}"

        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Билет #{ticket_id}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Arial, sans-serif;
            background: linear-gradient(135deg, #2b0b3f, #1a0525);
            min-height: 100vh;
            display: flex;
            justify-content: center;
            align-items: center;
            margin: 0;
            padding: 20px;
        }}
        .ticket {{
            background: white;
            max-width: 600px;
            width: 100%;
            border-radius: 24px;
            overflow: hidden;
            box-shadow: 0 20px 40px rgba(0,0,0,0.3);
        }}
        .header {{
            background: #6A11CB;
            color: white;
            padding: 30px;
            text-align: center;
        }}
        .header h1 {{ margin: 0; font-size: 2.5rem; }}
        .content {{ padding: 30px; }}
        .seat-info {{
            background: linear-gradient(135deg, #6A11CB, #2575FC);
            color: white;
            padding: 25px;
            text-align: center;
            border-radius: 16px;
            margin-bottom: 30px;
        }}
        .seat-info h2 {{ margin: 0; font-size: 2.5rem; }}
        .info-grid {{
            display: grid;
            grid-template-columns: 120px 1fr;
            gap: 15px;
            margin-bottom: 30px;
        }}
        .info-label {{ font-weight: bold; color: #666; }}
        .qr-section {{
            text-align: center;
            margin: 30px 0;
            padding: 20px;
            background: #f5f5f5;
            border-radius: 12px;
        }}
        .qr-code {{
            font-family: 'Courier New', monospace;
            font-size: 1.5rem;
            letter-spacing: 5px;
            background: white;
            padding: 15px;
            border-radius: 8px;
            display: inline-block;
        }}
        .footer {{
            text-align: center;
            padding: 20px;
            background: #f9f9f9;
            color: #999;
            font-size: 0.9rem;
        }}
        .btn-print {{
            background: #6A11CB;
            color: white;
            padding: 15px;
            border: none;
            border-radius: 8px;
            font-size: 1.1rem;
            cursor: pointer;
            width: 100%;
            margin-top: 20px;
        }}
        @media print {{
            body {{ background: white; padding: 0; }}
            .btn-print {{ display: none; }}
        }}
    </style>
</head>
<body>
    <div class="ticket">
        <div class="header">
            <h1>🎬 WebCinema</h1>
            <p>Электронный билет #{ticket_id}</p>
        </div>
        <div class="content">
            <div class="seat-info">
                <h2>Ряд {ticket['row_number']} / Место {ticket['seat_number']}</h2>
            </div>
            <div class="info-grid">
                <div class="info-label">Фильм:</div><div>{ticket['movie']}</div>
                <div class="info-label">Дата:</div><div>{date_str}</div>
                <div class="info-label">Время:</div><div>{time_str}</div>
                <div class="info-label">Зал:</div><div>{ticket['hall']}</div>
                <div class="info-label">Зритель:</div><div>{ticket['customer_name']}</div>
                <div class="info-label">Цена:</div><div>{ticket['price']} ₽</div>
                <div class="info-label">Оплата:</div><div>{ticket['payment_method']}</div>
            </div>
            <div class="qr-section">
                <div class="qr-code">✦ {ticket_id:08d} ✦</div>
            </div>
            <button onclick="window.print()" class="btn-print">🖨️ Распечатать</button>
        </div>
        <div class="footer">Билет действителен только на указанный сеанс</div>
    </div>
</body>
</html>"""
        return HTMLResponse(content=html)

    except Exception as e:
        logger.error(f"Ошибка: {e}")
        return HTMLResponse(f"Ошибка: {e}", status_code=500)


# ========== ТЕСТ ==========

@app.get("/test")
async def test():
    return {"message": "Hello World"}


if __name__ == "__main__":
    uvicorn.run(
        "main:app",
        host="127.0.0.1",
        port=8000,
        reload=True
    )